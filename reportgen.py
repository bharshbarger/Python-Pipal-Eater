#!/usr/bin/env python
"""module to generate a docx report"""

#https://python-docx.readthedocs.io/en/latest/user/text.html
#https://python-docx.readthedocs.io/en/latest/user/quickstart.html
import time
import docx
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches
#from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

class Reportgen(object):
    """class to generate a docx report from module output"""
    def __init__(self):
        self.today = time.strftime("%m/%d/%Y")

    def run(self,\
        total,\
        unique,\
        top_10,\
        top_10_base,\
        lengths,\
        counts,\
        one_to_six,\
        trailing_number,\
        last_1digit,\
        last_2digit,\
        last_3digit,\
        last_4digit,\
        last_5digit,\
        charset):

        print('\n[+] Generating Appendix')

        #start docx object
        self.document = docx.Document()

        #add paragraph
        paragraph = self.document.add_paragraph()

        #first paragraph is boilerplate explaining this section
        run_paragraph = paragraph.add_run('\nThis Appendix Contains Analysis of Cracked Passwords.\n')
        font = run_paragraph.font
        font.name = 'Hind'
        font.size = Pt(11)

        #create total passwords and unique table
        total_table = self.document.add_table(rows=0, cols=1)
        #one line per cell
        total_table.style = 'Light Grid'
        #set table font
        font = total_table.style.font
        font.name = 'Hind'
        font.size = Pt(11)
        font.bold = False
        #try to do autofit, doesnt seem to be working in libreoffice, but autofit is default anyway....
        #possibly related https://github.com/python-openxml/python-docx/issues/209
        #total_table.allow_autofit
        cells = total_table.add_row().cells
        cells[0].text = str(total)
        cells = total_table.add_row().cells
        cells[0].text = str(unique)
        #hack to fix width? https://github.com/python-openxml/python-docx/issues/315
        '''for r in total_table.rows:
            for c in r._tr.tc_lst:
                tcW = c.tcPr.tcW
                tcW.type = 'auto'
                tcW.w = 0'''
        paragraph = self.document.add_paragraph()

        #second table - top 10 overall pws
        top_10_table = self.document.add_table(rows=0, cols=1)
        top_10_table.style = 'Light Grid'
        #set table font
        font = top_10_table.style.font
        font.name = 'Hind'
        font.size = Pt(11)
        for i, val in enumerate(top_10):
            cells = top_10_table.add_row().cells
            cells[0].text = str(top_10[i])
        paragraph = self.document.add_paragraph()

        #third table - top 10 overall base dict words
        top_10_base_table = self.document.add_table(rows=0, cols=1)
        top_10_base_table.style = 'Light Grid'
        #set table font
        font = top_10_base_table.style.font
        font.name = 'Hind'
        font.size = Pt(11)
        for i, val in enumerate(top_10_base):
            cells = top_10_base_table.add_row().cells
            cells[0].text = str(top_10_base[i])
        paragraph = self.document.add_paragraph()

        #fourth table - passwords by length
        length_order_table = self.document.add_table(rows=0, cols=1)
        length_order_table.style = 'Light Grid'
        #set table font
        font = length_order_table.style.font
        font.name = 'Hind'
        font.size = Pt(11)
        for i, val in enumerate(lengths):
            cells = length_order_table.add_row().cells
            cells[0].text = str(lengths[i])
        paragraph = self.document.add_paragraph()

        #fifth table - passwords by count 
        count_order_table = self.document.add_table(rows=0, cols=1)
        count_order_table.style = 'Light Grid'
        #set table font
        font = count_order_table.style.font
        font.name = 'Hind'
        font.size = Pt(11)
        for i, val in enumerate(counts):
            cells = count_order_table.add_row().cells
            cells[0].text = str(counts[i])
        paragraph = self.document.add_paragraph()

        #sixth table
        one_to_six_table = self.document.add_table(rows=0, cols=1)
        one_to_six_table.style = 'Light Grid'
        #set table font
        font = one_to_six_table.style.font
        font.name = 'Hind'
        font.size = Pt(11)
        for i, val in enumerate(one_to_six):
            cells = one_to_six_table.add_row().cells
            cells[0].text = str(one_to_six[i])
        paragraph = self.document.add_paragraph()

        #seventh table - the trailing number
        trailing_number_table = self.document.add_table(rows=0, cols=1)
        trailing_number_table.style = 'Light Grid'
        #set table font
        font = trailing_number_table.style.font
        font.name = 'Hind'
        font.size = Pt(11)
        for i, val in enumerate(trailing_number):
            cells = trailing_number_table.add_row().cells
            cells[0].text = str(trailing_number[i])
        paragraph = self.document.add_paragraph()

        #eigth table - last digit in pw
        last_1digit_table = self.document.add_table(rows=0, cols=1)
        last_1digit_table.style = 'Light Grid'
        #set table font
        font = last_1digit_table.style.font
        font.name = 'Hind'
        font.size = Pt(11)
        for i, val in enumerate(last_1digit):
            cells = last_1digit_table.add_row().cells
            cells[0].text = str(last_1digit[i])
        paragraph = self.document.add_paragraph()

        #ninth table - last 2 digits of pw
        last_2digit_table = self.document.add_table(rows=0, cols=1)
        last_2digit_table.style = 'Light Grid'
        #set table font
        font = last_2digit_table.style.font
        font.name = 'Hind'
        font.size = Pt(11)
        for i, val in enumerate(last_2digit):
            cells = last_2digit_table.add_row().cells
            cells[0].text = str(last_2digit[i])
        paragraph = self.document.add_paragraph()

        #tenth table - last 3 digits of pw
        last_3digit_table = self.document.add_table(rows=0, cols=1)
        last_3digit_table.style = 'Light Grid'
        #set table font
        font = last_3digit_table.style.font
        font.name = 'Hind'
        font.size = Pt(11)
        for i, val in enumerate(last_3digit):
            cells = last_3digit_table.add_row().cells
            cells[0].text = str(last_3digit[i])
        paragraph = self.document.add_paragraph()

        #eleventh table - last 4 digits of pw
        last_4digit_table = self.document.add_table(rows=0, cols=1)
        last_4digit_table.style = 'Light Grid'
        #set table font
        font = last_4digit_table.style.font
        font.name = 'Hind'
        font.size = Pt(11)
        for i, val in enumerate(last_4digit):
            cells = last_4digit_table.add_row().cells
            cells[0].text = str(last_4digit[i])
        paragraph = self.document.add_paragraph()

        #twelvth table - last 5 digits of pw
        last_5digit_table = self.document.add_table(rows=0, cols=1)
        last_5digit_table.style = 'Light Grid'
        #set table font
        font = last_5digit_table.style.font
        font.name = 'Hind'
        font.size = Pt(11)
        for i, val in enumerate(last_5digit):
            cells = last_5digit_table.add_row().cells
            cells[0].text = str(last_5digit[i])
        paragraph = self.document.add_paragraph()

        #thirteenth table - characterset block
        charset_table = self.document.add_table(rows=0, cols=1)
        charset_table.style = 'Light Grid'
        #set table font
        font = charset_table.style.font
        font.name = 'Hind'
        font.size = Pt(11)
        for i, val in enumerate(charset):
            cells = charset_table.add_row().cells
            cells[0].text = str(charset[i])
   

        #have a look-see at all styles in dist-packages/docx/templates/default-styles.xml
        # you need to pull the grid styles out of that xml and into a file called docx_grid_styles.txt
        '''with open('docx_grid_styles.txt') as f:
            grid_styles = (f.readlines())
            grid_styles = [x.strip() for x in grid_styles]

        for g in grid_styles:
            paragraph = self.document.add_paragraph()
            run_paragraph = paragraph.add_run(g)

            charset_table = self.document.add_table(rows=0, cols=1)
            
            try:
                charset_table.style = g
            except:
                continue
            #set table font

            font = charset_table.style.font
            font.name = 'Hind'
            font.size = Pt(11)
            for i, val in enumerate(charset):
                cells = charset_table.add_row().cells
                cells[0].text = str(charset[i])
            paragraph = self.document.add_paragraph()'''

        self.document.save('Pipal_Appendix.docx'.format())
        print('[+] Appendix Generated! Yay!')
        print('''   +      o     +              o   
    +             o     +       +
    o          +
        o  +           +        +
        +        o     o       +        o
        -_-_-_-_-_-_-_,------,      o 
        _-_-_-_-_-_-_-|   /\_/\  
        -_-_-_-_-_-_-~|__( ^ .^)  +     +  
        _-_-_-_-_-_-_-""  ""      
        +      o         o   +       o
            +         +
            o        o         o      o     +
                o           +
                +      +     o        o      +''')